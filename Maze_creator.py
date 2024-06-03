import random
from docx import Document
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os


class Console:

    def __init__(self):
        self.sessions = []
        
        self.sessions.append(Session())
        while True:
            print("Maze creator")
            print("type -i for info")
            self.GetInput()
            os.system("cls")


    def PrintInfo(self):
        print("Info")
        print()
        print("type -num if you want to set the amount of mazes. They all will be generated differently")
        print()
        print("type -size to set height and width of your maze. Each of dimensions has to have at least 2 but not bigger than 30 columns/rows")
        print()
        print("type -create to generate a maze. O stands for a wall")
        print()
        print("type -cf to create a Word file with all last saved mazes")
        print()
        print("type -status to see last saved mazes and their amount")

    def PrintStatus(self):
        if(len(self.sessions)<2):
            print("no mazes are creating")
        else:
            print(str(self.sessions[len(self.sessions)-2].num) + " maze(-s) is/are created")
            print()
            self.sessions[len(self.sessions)-2].ShowAllMazes()

    def GetInput(self):

        inp = input()
        if inp == "-i":
            self.PrintInfo()
        elif inp == "-create":
            self.sessions[len(self.sessions)-1].CreateMazes()
            self.sessions[len(self.sessions)-1].ShowAllMazes()
            self.sessions.append(Session())
        elif inp == "-num":
            self.sessions[len(self.sessions)-1].getNumOfMazes()
        elif inp == "-size":
            self.sessions[len(self.sessions)-1].GetWidth()
            self.sessions[len(self.sessions)-1].GetHeight()
        elif inp == "-cf":
            FileName = input("Enter the name of the file ")
            self.sessions[len(self.sessions)-2].SaveAllMazesInFile(FileName)
            print("Succesfully!")
        elif inp == "-status":
            self.PrintStatus()
        else:
            print("the command " + inp + " is not defined")
        print()
        input("Press -Enter to continue")

class Session:
    def __init__(self):
        self.num = 1
        self.__allMazes = []
        self.Width = 3
        self.Height = 3
        self.__numOfLoadedMazes = 0
        self.Seed = [1, 1]
    
    def GetWidth(self):
        while True:
            marker = True
            try:
                Width = int(input("Enter the width of your Maze: "))
            except:
                print("Invalide input. Please, try again")
                marker = False
            if not marker :
                pass
            elif Width < 2 or Width > 30:
                print("the maze must have the width of at least 2 columns and max. 30")
            else:
                break
        self.Width = Width

    def GetHeight(self):
        while True:
            marker = True
            try:
                Height = int(input("Enter the height of your Maze: "))
            except:
                print("Invalide input. Please, try again")
                marker = False
            if not marker :
                pass
            elif Height < 2 or Height > 30:
                print("the maze must have the height of at least 2 rows and max. 30")
            else:
                break
        self.Height = Height

    def getNumOfMazes(self):
        while True:
            marker = True
            try:
                num = int(input("How much mazes do you want to generate? "))
            except:
                print("Invalide input. Please, try again")
                marker = False
            if not marker :
                pass
            elif num < 1:
                print("at least 1 maze, please")
            else:
                break
        self.num = num

    def CreateMazes(self):
        for i in range(self.num):
            M = Maze(self.Width, self.Height, self.Seed)
            M.Create()
            self.__allMazes.append(M)



    def ShowAllMazes(self):
        line = ""
        for i in self.__allMazes:
            for j in i.FinalMaze:
                for k in j:
                    line += k
                print(line)
                line = ""
            print("")
        
    def SaveAllMazesInFile(self, name):
        doc = Document()
        doc.add_heading("Mazes", 0)
        for i in self.__allMazes:
            self.CreateMazeFileMSWord(i, name, doc)
            self.__numOfLoadedMazes += 1


    def CreateMazeFileMSWord(self, Maze, fileName, doc):
        doc.add_paragraph(str(self.__numOfLoadedMazes))
        table = doc.add_table(rows=len(Maze.FinalMaze), cols = len(Maze.FinalMaze[0]))
        table.style = 'Table Grid'
        for i in range(len(Maze.FinalMaze)):
            row = table.rows[i].cells
            percent = (float(i)/float(len(Maze.FinalMaze)))*100
            self.TickLoader(percent)
            for j in range(len(Maze.FinalMaze[i])):
                if Maze.FinalMaze[i][j] != " ":
                    shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
                    table.cell(i, j)._tc.get_or_add_tcPr().append(shading_elm)

        doc.save(fileName+".docx")


    def TickLoader(self, percent):
        os.system("cls")
        print("Loading...")
        print(str(self.__numOfLoadedMazes + 1) + "/" + str(self.num) + " mazes " + str(percent) + "%")



class Maze:

    def __init__(self, width, height, seed):
        self.__width = width
        self.__height = height
        self.__slashs = []
        self.__checkedSlashs = []
        self.__isRowFullyCheckedList = []
        self.__slashsOutput = []
        self.__boundedOutput = []
        self.FinalMaze = []
        self.__seed = seed

        self.GenerateLists()


    def GenerateLists(self):
        for i in range(self.__height + 2):
            self.__slashs.append([])
            for j in range(self.__width + 2):
                self.__slashs[i].append(" ")
        for i in range(self.__width):
            self.__isRowFullyCheckedList.append(False)
            self.__checkedSlashs.append([])
            for j in range(self.__width):
                self.__checkedSlashs[i].append(False)
        for i in range(0, len(self.__seed), 2):
            self.__slashs[self.__seed[i+1]][self.__seed[i]] = "/"
            self.__checkedSlashs[self.__seed[i+1]-1][self.__seed[i]-1] = True


    def Create(self):
    
        self.FillInitiateMaze()

        for i in range(self.__height):
            self.__slashsOutput.append("")
            for j in range(self.__width):
                if self.__slashs[i+1][j+1] == " ":
                    self.__slashsOutput[i] += " "
                else:
                    self.__slashsOutput[i] += "O"

        for i in range(self.__height + 1):
            self.__boundedOutput.append([])
            for j in range(self.__width + 1):
                if i == 0 or i == self.__height or j == 0 or j == self.__width:
                    self.__boundedOutput[i].append("B")
                else:
                    self.__boundedOutput[i].append(self.__slashsOutput[i][j])

        self.FinalMaze = self.StartWallBreaker(self.__boundedOutput)
        self.FinalMaze = self.StartWallBreaker(self.FinalMaze)
        self.FinalMaze = self.StartWallBreaker(self.FinalMaze)


        for i in range(len(self.FinalMaze)):
            for j in range(len(self.FinalMaze[i])):
                if self.FinalMaze[i][j] == "B":
                    self.FinalMaze[i][j] = "O"
                elif self.FinalMaze[i][j] == "E":
                    self.FinalMaze[i][j] = " "
                elif self.FinalMaze[i][j] == "O":
                    pass
                else:
                    self.FinalMaze[i][j] = " "
        return self.FinalMaze

    def FillInitiateMaze(self):
        isToReturn = True
        for i in self.__isRowFullyCheckedList:
            if i == False:
                isToReturn = False  
        if isToReturn:
            return
        pos = self.getRandomPosition()
        if len(pos) == 0:
            return
        state = random.randrange(0, 2)    
        self.check(pos[1], pos[0], state)

        self.FillInitiateMaze()


    def check(self, posx, posy, state):
        if state==1:
            if self.checkStraightSlash(posx, posy) == False:
                self.checkBackSlash(posx, posy)
        else:
            if self.checkBackSlash(posx, posy) == False:
                self.checkStraightSlash(posx, posy)


    def checkBackSlash(self, posx, posy):
        if self.__slashs[posx][posy]==" ":
            if self.__slashs[posx-1][posy-1] != "\\" and self.__slashs[posx+1][posy+1] != "\\":
                if self.__slashs[posx][posy+1] != "/" and self.__slashs[posx][posy-1] != "/" and self.__slashs[posx+1][posy] != "/" and self.__slashs[posx-1][posy] != "/":
                    self.__slashs[posx][posy] = "\\"
                    self.__checkedSlashs[posx - 1][posy - 1] = True
                    return True
                else:
                    self.__checkedSlashs[posx - 1][posy - 1] = True
                    return False
            else:
                self.__checkedSlashs[posx - 1][posy - 1] = True
                return False
        return False

    def checkStraightSlash(self, posx, posy):
        if self.__slashs[posx][posy]==" ":
            if self.__slashs[posx-1][posy+1] != "/" and self.__slashs[posx+1][posy-1] != "/":
                if self.__slashs[posx][posy+1] != "\\" and self.__slashs[posx][posy-1] != "\\" and self.__slashs[posx+1][posy] != "\\" and self.__slashs[posx-1][posy] != "\\":
                    self.__slashs[posx][posy] = "/"
                    self.__checkedSlashs[posx - 1][posy - 1] = True
                    return True
                else:
                    self.__checkedSlashs[posx - 1][posy - 1] = True
                    return False
            else:
                self.__checkedSlashs[posx - 1][posy - 1] = True
                return False
        return False

                
    def StartWallBreaker(self, generatedMaze): 
        numOfIteration = 0
        FreeMaze = generatedMaze
        for i in range(len(generatedMaze)):
            for j in range(len(generatedMaze[i])):
                if generatedMaze[i][j] == " ":
                    FreeMaze = self.MakeFreeClosedSpace(i, j, FreeMaze, numOfIteration, True)
                    numOfIteration += 1
        return FreeMaze

    def FindNeighbours(self, row, unit, symbol, Maze):
        if Maze[row - 1][unit] == symbol:
            return [row - 1, unit]
        
        elif Maze[row + 1][unit] == symbol:
            return [row + 1,unit]
        
        elif Maze[row][unit - 1] == symbol:
            return [row, unit - 1]
        
        elif Maze[row][unit + 1] == symbol:
            return [row, unit + 1]
        else:
            return []

    def CountTheNeighbouredWalls(self, row, unit, Maze):
        num = 0
        for i in range(row - 1, row + 2):
            for j in range(unit - 1, unit + 2):
                if Maze[i][j] == "O":
                    num+=1
        return num

    def FindTheMostCloseWall(self, row, unit, Maze):
        theWall = []
        counted = 0
        if Maze[row-1][unit] == "O":
            if self.CountTheNeighbouredWalls(row - 1, unit, Maze) > counted:
                counted = self.CountTheNeighbouredWalls(row - 1, unit, Maze)
                theWall = [row-1, unit]
        if Maze[row+1][unit] == "O":
            if self.CountTheNeighbouredWalls(row + 1, unit, Maze) > counted:
                counted = self.CountTheNeighbouredWalls(row + 1, unit, Maze)
                theWall = [row+1, unit]
        if Maze[row][unit - 1] == "O":
            if self.CountTheNeighbouredWalls(row, unit - 1, Maze) > counted:
                counted = self.CountTheNeighbouredWalls(row, unit - 1, Maze)
                theWall = [row, unit - 1]
        if Maze[row][unit + 1] == "O":
            if self.CountTheNeighbouredWalls(row, unit + 1, Maze) > counted:
                counted = self.CountTheNeighbouredWalls(row, unit + 1, Maze)
                theWall = [row, unit + 1]
        return theWall

    def MakeFreeClosedSpace(self, row, unit, Maze, num, isToBreak):
        MazeToReturn = Maze
        MazeToReturn[row][unit] = str(num)

        
        EmptyNeighbour = self.FindNeighbours(row, unit, " ", Maze)
        
        if EmptyNeighbour != []:
            if MazeToReturn[row-1][unit] == " ":
                MazeToReturn = self.MakeFreeClosedSpace(row-1, unit, MazeToReturn, num, isToBreak)
                isToBreak = False
            if MazeToReturn[row+1][unit] == " ":
                MazeToReturn = self.MakeFreeClosedSpace(row+1, unit, MazeToReturn, num, isToBreak)
                isToBreak = False
            if MazeToReturn[row][unit - 1] == " ":
                MazeToReturn = self.MakeFreeClosedSpace(row, unit-1, MazeToReturn, num, isToBreak)
                isToBreak = False
            if MazeToReturn[row][unit + 1] == " ":
                MazeToReturn = self.MakeFreeClosedSpace(row, unit+1, MazeToReturn, num, isToBreak)
                isToBreak = False
        else:
            
            for i in range(num):
                numberedNeighbours = self.FindNeighbours(row, unit, str(i), MazeToReturn)
                if numberedNeighbours != []:
                    return MazeToReturn
            BoundsNeighbours = self.FindNeighbours(row, unit, "B", MazeToReturn)

            if BoundsNeighbours != []:
                MazeToReturn[BoundsNeighbours[0]][BoundsNeighbours[1]]="E"
                return MazeToReturn
            
            if self.FindNeighbours(row, unit, "O", MazeToReturn) != [] and isToBreak == True:
                toBreak = self.FindTheMostCloseWall(row, unit, MazeToReturn)
                MazeToReturn[toBreak[0]][toBreak[1]] = str(num)
                return MazeToReturn
            
        return MazeToReturn

    def getRandomPosition(self):
        ListOfRows = []
        for i in range(len(self.__isRowFullyCheckedList)):
            if self.__isRowFullyCheckedList[i]==False:
                ListOfRows.append(i)
        if len(ListOfRows) == 0:
            return []
        y = random.choice(ListOfRows)
        ListOfDegrees = []
        for i in range(len(self.__checkedSlashs[y])):
            if self.__checkedSlashs[y][i] == False:
                ListOfDegrees.append(i)
        if len(ListOfDegrees) == 0:
            self.__isRowFullyCheckedList[y] = True
            self.getRandomPosition()
            return self.getRandomPosition()
        x = random.choice(ListOfDegrees)
        res = [x + 1, y + 1]
        return res
        
random.seed(None, 2)
console = Console()







